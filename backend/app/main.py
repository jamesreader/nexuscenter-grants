"""
Nexus Center Grants Management System - API

AI-powered grant discovery and application management for nonprofits.
"""

import logging
import os
from contextlib import asynccontextmanager
from decimal import Decimal
from typing import Optional
from uuid import UUID

import httpx
from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sqlalchemy import select, func, any_
from sqlalchemy.ext.asyncio import AsyncSession

from app.db import get_db, init_db, close_db
from app.models import GrantOpportunity, Application, GrantStatus, ApplicationStatus

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

__version__ = "1.0.0"

CORTEX_URL = os.getenv("CORTEX_URL", "http://ai-gateway.cortex.svc.cluster.local:8000")


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan handler."""
    logger.info("Starting Nexus Center Grants API...")
    await init_db()
    yield
    logger.info("Shutting down Nexus Center Grants API...")
    await close_db()


app = FastAPI(
    title="Nexus Center Grants Management",
    version=__version__,
    description="AI-powered grant discovery and management for Nexus Center for IDD Care",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================================
# Health & Stats
# ============================================================================

@app.get("/health")
async def health():
    return {"status": "healthy", "version": __version__}


@app.get("/api/v1/stats")
async def get_stats(db: AsyncSession = Depends(get_db)):
    """Dashboard statistics."""
    # Count grants by status
    grants_open = await db.scalar(
        select(func.count(GrantOpportunity.id))
        .where(GrantOpportunity.status == GrantStatus.OPEN)
    )
    grants_total = await db.scalar(select(func.count(GrantOpportunity.id)))

    # Count applications by status
    apps_in_progress = await db.scalar(
        select(func.count(Application.id))
        .where(Application.status.in_([
            ApplicationStatus.DRAFT,
            ApplicationStatus.IN_PROGRESS,
            ApplicationStatus.READY_FOR_REVIEW
        ]))
    )
    apps_awarded = await db.scalar(
        select(func.count(Application.id))
        .where(Application.status == ApplicationStatus.AWARDED)
    )

    # Sum awarded amounts
    total_awarded = await db.scalar(
        select(func.sum(Application.awarded_amount))
        .where(Application.status == ApplicationStatus.AWARDED)
    ) or 0

    # Upcoming deadlines (next 30 days)
    from datetime import datetime, timedelta
    deadline_cutoff = datetime.now() + timedelta(days=30)
    upcoming = await db.scalar(
        select(func.count(GrantOpportunity.id))
        .where(GrantOpportunity.status == GrantStatus.OPEN)
        .where(GrantOpportunity.deadline <= deadline_cutoff)
    )

    return {
        "grants_open": grants_open or 0,
        "grants_total": grants_total or 0,
        "applications_in_progress": apps_in_progress or 0,
        "applications_awarded": apps_awarded or 0,
        "total_awarded": float(total_awarded),
        "upcoming_deadlines": upcoming or 0,
    }


# ============================================================================
# Grant Discovery (Cortex Integration)
# ============================================================================

class DiscoverRequest(BaseModel):
    """Request to discover relevant grants."""
    query: str
    include_dms: bool = True
    max_results: int = 10


class DiscoverResponse(BaseModel):
    """Response from grant discovery."""
    message: str
    sources: list[dict]
    processing_time_ms: float


@app.post("/api/v1/discover", response_model=DiscoverResponse)
async def discover_grants(request: DiscoverRequest):
    """
    Search for relevant grants using Cortex RAG.

    Searches both the grants database and DMS for supporting documents.
    """
    sources = ["grants"]
    if request.include_dms:
        sources.append("dms")

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{CORTEX_URL}/v1/rag",
                json={
                    "message": request.query,
                    "sources": sources,
                    "scope": "internal",
                    "max_iterations": 4,
                }
            )
            response.raise_for_status()
            data = response.json()

            return DiscoverResponse(
                message=data.get("message", ""),
                sources=data.get("sources", []),
                processing_time_ms=data.get("processing_time_ms", 0),
            )

    except httpx.HTTPStatusError as e:
        logger.error(f"Cortex error: {e.response.text}")
        raise HTTPException(502, "AI service unavailable")
    except httpx.RequestError as e:
        logger.error(f"Cortex connection error: {e}")
        raise HTTPException(503, "AI service connection failed")


# ============================================================================
# Grant Opportunities CRUD
# ============================================================================

class GrantCreate(BaseModel):
    """Create a new grant opportunity."""
    title: str
    agency: Optional[str] = None
    description: Optional[str] = None
    funding_amount_min: Optional[float] = None
    funding_amount_max: Optional[float] = None
    deadline: Optional[str] = None
    eligibility: Optional[str] = None
    categories: Optional[list[str]] = None
    source_url: Optional[str] = None


class GrantResponse(BaseModel):
    """Grant opportunity response."""
    id: str
    title: str
    agency: Optional[str]
    description: Optional[str]
    funding_amount_min: Optional[float]
    funding_amount_max: Optional[float]
    deadline: Optional[str]
    status: str
    categories: Optional[list[str]]

    class Config:
        from_attributes = True


@app.get("/api/v1/grants")
async def list_grants(
    status: Optional[str] = Query(None, description="Filter by status"),
    sort_by: str = Query("relevance", description="Sort by: relevance, deadline, created_at"),
    category: Optional[str] = Query(None, description="Filter by category"),
    min_score: Optional[float] = Query(None, ge=0, le=1, description="Minimum relevance score"),
    include_dismissed: bool = Query(False, description="Include dismissed grants"),
    limit: int = Query(20, le=100),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """List grant opportunities with relevance scores."""
    query = select(GrantOpportunity)

    # Filter by status
    if status:
        try:
            status_enum = GrantStatus(status)
            query = query.where(GrantOpportunity.status == status_enum)
        except ValueError:
            raise HTTPException(400, f"Invalid status: {status}")

    # Filter by category (use any_ for PostgreSQL array contains)
    if category:
        query = query.where(category == any_(GrantOpportunity.categories))

    # Filter by minimum relevance score
    if min_score is not None:
        query = query.where(GrantOpportunity.relevance_score >= min_score)

    # Exclude dismissed unless requested
    if not include_dismissed:
        query = query.where(GrantOpportunity.is_dismissed == False)

    # Sort order
    if sort_by == "relevance":
        query = query.order_by(GrantOpportunity.relevance_score.desc().nullslast())
    elif sort_by == "deadline":
        query = query.order_by(GrantOpportunity.deadline.asc().nullslast())
    elif sort_by == "created_at":
        query = query.order_by(GrantOpportunity.created_at.desc())
    else:
        query = query.order_by(GrantOpportunity.relevance_score.desc().nullslast())

    query = query.offset(offset).limit(limit)
    result = await db.execute(query)
    grants = result.scalars().all()

    return {
        "grants": [
            {
                "id": str(g.id),
                "title": g.title,
                "agency": g.agency,
                "description": (g.description or "")[:200],
                "funding_amount_min": float(g.funding_amount_min) if g.funding_amount_min else None,
                "funding_amount_max": float(g.funding_amount_max) if g.funding_amount_max else None,
                "deadline": g.deadline.isoformat() if g.deadline else None,
                "status": g.status.value,
                "categories": g.categories,
                "relevance_score": float(g.relevance_score) if g.relevance_score else None,
                "relevance_notes": g.relevance_notes,
                "is_flagged": g.is_flagged,
                "is_dismissed": g.is_dismissed,
            }
            for g in grants
        ],
        "total": len(grants),
        "offset": offset,
        "limit": limit,
    }


@app.post("/api/v1/grants")
async def create_grant(
    grant: GrantCreate,
    db: AsyncSession = Depends(get_db),
):
    """Create a new grant opportunity (manual entry)."""
    from datetime import datetime

    new_grant = GrantOpportunity(
        title=grant.title,
        agency=grant.agency,
        description=grant.description,
        funding_amount_min=grant.funding_amount_min,
        funding_amount_max=grant.funding_amount_max,
        deadline=datetime.fromisoformat(grant.deadline) if grant.deadline else None,
        eligibility=grant.eligibility,
        categories=grant.categories,
        source_url=grant.source_url,
        content_text=f"{grant.title}\n{grant.description or ''}\n{grant.eligibility or ''}",
    )
    db.add(new_grant)
    await db.commit()
    await db.refresh(new_grant)

    return {"id": str(new_grant.id), "status": "created"}


@app.get("/api/v1/grants/flagged")
async def list_flagged_grants(
    limit: int = Query(50, le=100),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """List all flagged grants."""
    query = (
        select(GrantOpportunity)
        .where(GrantOpportunity.is_flagged == True)
        .order_by(GrantOpportunity.deadline.asc().nullslast())
        .offset(offset)
        .limit(limit)
    )
    result = await db.execute(query)
    grants = result.scalars().all()

    return {
        "grants": [
            {
                "id": str(g.id),
                "title": g.title,
                "agency": g.agency,
                "deadline": g.deadline.isoformat() if g.deadline else None,
                "relevance_score": float(g.relevance_score) if g.relevance_score else None,
                "user_notes": g.user_notes,
            }
            for g in grants
        ],
        "total": len(grants),
    }


@app.get("/api/v1/grants/{grant_id}")
async def get_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Get a specific grant opportunity."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    return {
        "id": str(grant.id),
        "title": grant.title,
        "agency": grant.agency,
        "description": grant.description,
        "funding_amount_min": float(grant.funding_amount_min) if grant.funding_amount_min else None,
        "funding_amount_max": float(grant.funding_amount_max) if grant.funding_amount_max else None,
        "deadline": grant.deadline.isoformat() if grant.deadline else None,
        "status": grant.status.value,
        "categories": grant.categories,
        "eligibility": grant.eligibility,
        "requirements": grant.requirements,
        "source": grant.source.value,
        "source_url": grant.source_url,
        "relevance_score": float(grant.relevance_score) if grant.relevance_score else None,
        "relevance_notes": grant.relevance_notes,
        "is_flagged": grant.is_flagged,
        "is_dismissed": grant.is_dismissed,
        "user_notes": grant.user_notes,
        "created_at": grant.created_at.isoformat(),
    }


# ============================================================================
# Grant Management (dismiss, flag, notes)
# ============================================================================

class GrantUpdateRequest(BaseModel):
    """Request to update a grant."""
    is_dismissed: Optional[bool] = None
    is_flagged: Optional[bool] = None
    user_notes: Optional[str] = None
    status: Optional[str] = None


@app.patch("/api/v1/grants/{grant_id}")
async def update_grant(
    grant_id: str,
    request: GrantUpdateRequest,
    db: AsyncSession = Depends(get_db),
):
    """Update a grant opportunity (dismiss, flag, add notes)."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    if request.is_dismissed is not None:
        grant.is_dismissed = request.is_dismissed
    if request.is_flagged is not None:
        grant.is_flagged = request.is_flagged
    if request.user_notes is not None:
        grant.user_notes = request.user_notes
    if request.status is not None:
        try:
            grant.status = GrantStatus(request.status)
        except ValueError:
            raise HTTPException(400, f"Invalid status: {request.status}")

    await db.commit()

    return {
        "id": str(grant.id),
        "is_dismissed": grant.is_dismissed,
        "is_flagged": grant.is_flagged,
        "user_notes": grant.user_notes,
        "status": grant.status.value,
    }


@app.post("/api/v1/grants/{grant_id}/dismiss")
async def dismiss_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Dismiss a grant (mark as not relevant)."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    grant.is_dismissed = True
    await db.commit()

    return {"id": str(grant.id), "is_dismissed": True}


@app.delete("/api/v1/grants/{grant_id}/dismiss")
async def undismiss_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Undo dismiss on a grant."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    grant.is_dismissed = False
    await db.commit()

    return {"id": str(grant.id), "is_dismissed": False}


@app.post("/api/v1/grants/{grant_id}/flag")
async def flag_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Flag a grant for follow-up."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    grant.is_flagged = True
    await db.commit()

    return {"id": str(grant.id), "is_flagged": True}


@app.delete("/api/v1/grants/{grant_id}/flag")
async def unflag_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Remove flag from a grant."""
    try:
        uuid = UUID(grant_id)
    except ValueError:
        raise HTTPException(400, "Invalid grant ID")

    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    grant = result.scalar_one_or_none()

    if not grant:
        raise HTTPException(404, "Grant not found")

    grant.is_flagged = False
    await db.commit()

    return {"id": str(grant.id), "is_flagged": False}


# ============================================================================
# Applications CRUD
# ============================================================================

@app.get("/api/v1/applications")
async def list_applications(
    status: Optional[str] = Query(None),
    limit: int = Query(20, le=100),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """List grant applications."""
    query = (
        select(Application)
        .order_by(Application.submission_deadline.asc())
    )

    if status:
        try:
            status_enum = ApplicationStatus(status)
            query = query.where(Application.status == status_enum)
        except ValueError:
            raise HTTPException(400, f"Invalid status: {status}")

    query = query.offset(offset).limit(limit)
    result = await db.execute(query)
    apps = result.scalars().all()

    return {
        "applications": [
            {
                "id": str(a.id),
                "opportunity_id": str(a.opportunity_id),
                "project_title": a.project_title,
                "status": a.status.value,
                "submission_deadline": a.submission_deadline.isoformat() if a.submission_deadline else None,
                "requested_amount": float(a.requested_amount) if a.requested_amount else None,
            }
            for a in apps
        ],
        "total": len(apps),
    }


@app.post("/api/v1/applications")
async def create_application(
    opportunity_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Start a new application for a grant opportunity."""
    try:
        uuid = UUID(opportunity_id)
    except ValueError:
        raise HTTPException(400, "Invalid opportunity ID")

    # Verify opportunity exists
    result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == uuid)
    )
    opportunity = result.scalar_one_or_none()
    if not opportunity:
        raise HTTPException(404, "Grant opportunity not found")

    # Create application
    new_app = Application(
        opportunity_id=uuid,
        submission_deadline=opportunity.deadline,
    )
    db.add(new_app)
    await db.commit()
    await db.refresh(new_app)

    return {"id": str(new_app.id), "status": "created"}


@app.get("/api/v1/applications/{application_id}")
async def get_application(
    application_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Get a specific application with full details."""
    try:
        uuid = UUID(application_id)
    except ValueError:
        raise HTTPException(400, "Invalid application ID")

    result = await db.execute(
        select(Application).where(Application.id == uuid)
    )
    app_obj = result.scalar_one_or_none()

    if not app_obj:
        raise HTTPException(404, "Application not found")

    # Get opportunity details
    opp_result = await db.execute(
        select(GrantOpportunity).where(GrantOpportunity.id == app_obj.opportunity_id)
    )
    opportunity = opp_result.scalar_one_or_none()

    return {
        "id": str(app_obj.id),
        "opportunity_id": str(app_obj.opportunity_id),
        "opportunity_title": opportunity.title if opportunity else None,
        "opportunity_agency": opportunity.agency if opportunity else None,
        "status": app_obj.status.value,
        "project_title": app_obj.project_title,
        "project_description": app_obj.project_description,
        "requested_amount": float(app_obj.requested_amount) if app_obj.requested_amount else None,
        "awarded_amount": float(app_obj.awarded_amount) if app_obj.awarded_amount else None,
        "local_match_amount": float(app_obj.local_match_amount) if app_obj.local_match_amount else None,
        "match_source": app_obj.match_source,
        "submission_deadline": app_obj.submission_deadline.isoformat() if app_obj.submission_deadline else None,
        "submitted_at": app_obj.submitted_at.isoformat() if app_obj.submitted_at else None,
        "decision_date": app_obj.decision_date.isoformat() if app_obj.decision_date else None,
        "narrative_draft": app_obj.narrative_draft,
        "budget_draft": app_obj.budget_draft,
        "requirements_checklist": app_obj.requirements_checklist,
        "supporting_documents": app_obj.supporting_documents,
        "notes": app_obj.notes,
        "history": app_obj.history,
        "assigned_to": app_obj.assigned_to,
        "confirmation_number": app_obj.confirmation_number,
        "submission_method": app_obj.submission_method,
        "created_at": app_obj.created_at.isoformat(),
        "updated_at": app_obj.updated_at.isoformat(),
    }


class ApplicationUpdateRequest(BaseModel):
    """Request to update an application."""
    project_title: Optional[str] = None
    project_description: Optional[str] = None
    requested_amount: Optional[float] = None
    local_match_amount: Optional[float] = None
    match_source: Optional[str] = None
    submission_deadline: Optional[str] = None
    narrative_draft: Optional[str] = None
    budget_draft: Optional[dict] = None
    requirements_checklist: Optional[dict] = None
    supporting_documents: Optional[list[dict]] = None
    notes: Optional[str] = None
    assigned_to: Optional[str] = None


@app.patch("/api/v1/applications/{application_id}")
async def update_application(
    application_id: str,
    request: ApplicationUpdateRequest,
    db: AsyncSession = Depends(get_db),
):
    """Update an application."""
    try:
        uuid = UUID(application_id)
    except ValueError:
        raise HTTPException(400, "Invalid application ID")

    result = await db.execute(
        select(Application).where(Application.id == uuid)
    )
    app_obj = result.scalar_one_or_none()

    if not app_obj:
        raise HTTPException(404, "Application not found")

    # Update fields if provided
    if request.project_title is not None:
        app_obj.project_title = request.project_title
    if request.project_description is not None:
        app_obj.project_description = request.project_description
    if request.requested_amount is not None:
        app_obj.requested_amount = Decimal(str(request.requested_amount))
    if request.local_match_amount is not None:
        app_obj.local_match_amount = Decimal(str(request.local_match_amount))
    if request.match_source is not None:
        app_obj.match_source = request.match_source
    if request.submission_deadline is not None:
        from datetime import datetime as dt
        app_obj.submission_deadline = dt.fromisoformat(request.submission_deadline)
    if request.narrative_draft is not None:
        app_obj.narrative_draft = request.narrative_draft
    if request.budget_draft is not None:
        app_obj.budget_draft = request.budget_draft
    if request.requirements_checklist is not None:
        app_obj.requirements_checklist = request.requirements_checklist
    if request.supporting_documents is not None:
        app_obj.supporting_documents = request.supporting_documents
    if request.notes is not None:
        app_obj.notes = request.notes
    if request.assigned_to is not None:
        app_obj.assigned_to = request.assigned_to

    await db.commit()

    return {"id": str(app_obj.id), "status": "updated"}


class StatusChangeRequest(BaseModel):
    """Request to change application status."""
    status: str
    note: Optional[str] = None
    awarded_amount: Optional[float] = None
    confirmation_number: Optional[str] = None
    submission_method: Optional[str] = None


@app.post("/api/v1/applications/{application_id}/status")
async def change_application_status(
    application_id: str,
    request: StatusChangeRequest,
    db: AsyncSession = Depends(get_db),
):
    """Change application status with history tracking."""
    from datetime import datetime as dt

    try:
        uuid = UUID(application_id)
    except ValueError:
        raise HTTPException(400, "Invalid application ID")

    result = await db.execute(
        select(Application).where(Application.id == uuid)
    )
    app_obj = result.scalar_one_or_none()

    if not app_obj:
        raise HTTPException(404, "Application not found")

    # Validate status
    try:
        new_status = ApplicationStatus(request.status)
    except ValueError:
        raise HTTPException(400, f"Invalid status: {request.status}")

    old_status = app_obj.status

    # Update status
    app_obj.status = new_status

    # Handle specific status transitions
    if new_status == ApplicationStatus.SUBMITTED:
        app_obj.submitted_at = dt.utcnow()
        if request.confirmation_number:
            app_obj.confirmation_number = request.confirmation_number
        if request.submission_method:
            app_obj.submission_method = request.submission_method

    if new_status in [ApplicationStatus.AWARDED, ApplicationStatus.REJECTED]:
        app_obj.decision_date = dt.utcnow()
        if request.awarded_amount:
            app_obj.awarded_amount = Decimal(str(request.awarded_amount))

    # Add to history
    history_entry = {
        "timestamp": dt.utcnow().isoformat(),
        "from_status": old_status.value,
        "to_status": new_status.value,
        "note": request.note,
    }
    if app_obj.history is None:
        app_obj.history = []
    app_obj.history = app_obj.history + [history_entry]

    await db.commit()

    return {
        "id": str(app_obj.id),
        "status": app_obj.status.value,
        "previous_status": old_status.value,
    }


@app.delete("/api/v1/applications/{application_id}")
async def withdraw_application(
    application_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Withdraw/delete an application."""
    from datetime import datetime as dt

    try:
        uuid = UUID(application_id)
    except ValueError:
        raise HTTPException(400, "Invalid application ID")

    result = await db.execute(
        select(Application).where(Application.id == uuid)
    )
    app_obj = result.scalar_one_or_none()

    if not app_obj:
        raise HTTPException(404, "Application not found")

    # If submitted, mark as withdrawn instead of deleting
    if app_obj.status in [ApplicationStatus.SUBMITTED, ApplicationStatus.UNDER_REVIEW]:
        app_obj.status = ApplicationStatus.WITHDRAWN
        history_entry = {
            "timestamp": dt.utcnow().isoformat(),
            "from_status": app_obj.status.value,
            "to_status": "withdrawn",
            "note": "Application withdrawn",
        }
        if app_obj.history is None:
            app_obj.history = []
        app_obj.history = app_obj.history + [history_entry]
        await db.commit()
        return {"id": str(app_obj.id), "status": "withdrawn"}
    else:
        # Delete draft applications
        await db.delete(app_obj)
        await db.commit()
        return {"id": application_id, "status": "deleted"}


# ============================================================================
# Reports
# ============================================================================

@app.get("/api/v1/reports/pipeline")
async def pipeline_report(db: AsyncSession = Depends(get_db)):
    """Application pipeline summary - grants and applications by status."""
    from datetime import datetime as dt, timedelta, timezone

    now = dt.now(timezone.utc)

    # Count grants by status
    grant_counts = {}
    for status in GrantStatus:
        count = await db.scalar(
            select(func.count(GrantOpportunity.id))
            .where(GrantOpportunity.status == status)
        )
        grant_counts[status.value] = count or 0

    # Count applications by status
    app_counts = {}
    for status in ApplicationStatus:
        count = await db.scalar(
            select(func.count(Application.id))
            .where(Application.status == status)
        )
        app_counts[status.value] = count or 0

    # High relevance grants (score > 0.7)
    high_relevance = await db.scalar(
        select(func.count(GrantOpportunity.id))
        .where(GrantOpportunity.relevance_score >= 0.7)
        .where(GrantOpportunity.status == GrantStatus.OPEN)
    )

    # Flagged grants count
    flagged = await db.scalar(
        select(func.count(GrantOpportunity.id))
        .where(GrantOpportunity.is_flagged == True)
    )

    # Upcoming deadlines (next 30 days)
    deadline_cutoff = now + timedelta(days=30)
    upcoming_deadlines = await db.scalar(
        select(func.count(GrantOpportunity.id))
        .where(GrantOpportunity.status == GrantStatus.OPEN)
        .where(GrantOpportunity.deadline <= deadline_cutoff)
        .where(GrantOpportunity.deadline >= now)
    )

    return {
        "grants": grant_counts,
        "applications": app_counts,
        "high_relevance_open": high_relevance or 0,
        "flagged": flagged or 0,
        "upcoming_deadlines_30d": upcoming_deadlines or 0,
    }


@app.get("/api/v1/reports/funding")
async def funding_report(db: AsyncSession = Depends(get_db)):
    """Funding summary - amounts requested, awarded, pending."""
    # Total requested (all non-withdrawn/rejected applications)
    total_requested = await db.scalar(
        select(func.sum(Application.requested_amount))
        .where(Application.status.notin_([
            ApplicationStatus.WITHDRAWN,
            ApplicationStatus.REJECTED,
        ]))
    )

    # Total awarded
    total_awarded = await db.scalar(
        select(func.sum(Application.awarded_amount))
        .where(Application.status == ApplicationStatus.AWARDED)
    )

    # Pending (submitted or under review)
    pending_count = await db.scalar(
        select(func.count(Application.id))
        .where(Application.status.in_([
            ApplicationStatus.SUBMITTED,
            ApplicationStatus.UNDER_REVIEW,
            ApplicationStatus.ADDITIONAL_INFO_REQUESTED,
        ]))
    )
    pending_amount = await db.scalar(
        select(func.sum(Application.requested_amount))
        .where(Application.status.in_([
            ApplicationStatus.SUBMITTED,
            ApplicationStatus.UNDER_REVIEW,
            ApplicationStatus.ADDITIONAL_INFO_REQUESTED,
        ]))
    )

    # In progress (draft, in_progress, ready_for_review)
    in_progress_count = await db.scalar(
        select(func.count(Application.id))
        .where(Application.status.in_([
            ApplicationStatus.DRAFT,
            ApplicationStatus.IN_PROGRESS,
            ApplicationStatus.READY_FOR_REVIEW,
        ]))
    )

    return {
        "total_requested": float(total_requested) if total_requested else 0,
        "total_awarded": float(total_awarded) if total_awarded else 0,
        "pending_applications": pending_count or 0,
        "pending_amount": float(pending_amount) if pending_amount else 0,
        "in_progress_applications": in_progress_count or 0,
    }


@app.get("/api/v1/reports/deadlines")
async def deadlines_report(
    days: int = Query(60, description="Number of days to look ahead"),
    db: AsyncSession = Depends(get_db),
):
    """Upcoming deadlines report."""
    from datetime import datetime as dt, timedelta, timezone

    now = dt.now(timezone.utc)
    cutoff = now + timedelta(days=days)

    result = await db.execute(
        select(GrantOpportunity)
        .where(GrantOpportunity.status == GrantStatus.OPEN)
        .where(GrantOpportunity.deadline >= now)
        .where(GrantOpportunity.deadline <= cutoff)
        .order_by(GrantOpportunity.deadline.asc())
        .limit(50)
    )
    grants = result.scalars().all()

    return {
        "period_days": days,
        "count": len(grants),
        "deadlines": [
            {
                "id": str(g.id),
                "title": g.title,
                "agency": g.agency,
                "deadline": g.deadline.isoformat() if g.deadline else None,
                "days_remaining": (g.deadline - now).days if g.deadline else None,
                "relevance_score": float(g.relevance_score) if g.relevance_score else None,
                "is_flagged": g.is_flagged,
                "funding_max": float(g.funding_amount_max) if g.funding_amount_max else None,
            }
            for g in grants
        ],
    }


@app.get("/api/v1/reports/categories")
async def categories_report(db: AsyncSession = Depends(get_db)):
    """Grants by category."""
    # Get all grants with categories
    result = await db.execute(
        select(GrantOpportunity)
        .where(GrantOpportunity.status == GrantStatus.OPEN)
    )
    grants = result.scalars().all()

    # Count by category
    category_counts: dict = {}
    for g in grants:
        if g.categories:
            for cat in g.categories:
                if cat not in category_counts:
                    category_counts[cat] = {"count": 0, "high_relevance": 0}
                category_counts[cat]["count"] += 1
                if g.relevance_score and g.relevance_score >= 0.7:
                    category_counts[cat]["high_relevance"] += 1

    return {
        "categories": [
            {
                "name": cat,
                "count": data["count"],
                "high_relevance": data["high_relevance"],
            }
            for cat, data in sorted(
                category_counts.items(),
                key=lambda x: x[1]["count"],
                reverse=True
            )
        ],
    }


# ============================================================================
# Grants Ingestion
# ============================================================================

class IngestRequest(BaseModel):
    """Request to trigger grants ingestion."""
    keywords: Optional[list[str]] = None
    categories: Optional[list[str]] = None
    max_results: int = 500


class IngestResponse(BaseModel):
    """Response from grants ingestion."""
    total_fetched: int
    new_grants: int
    updated_grants: int
    details_fetched: int
    scored: int
    errors: list[str]


@app.post("/api/v1/ingest/grants-gov", response_model=IngestResponse)
async def ingest_from_grants_gov(
    request: IngestRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Ingest grants from grants.gov.

    Fetches federal grant opportunities relevant to nonprofit organizations.
    Can filter by keywords or funding categories.
    """
    from app.ingest.grants_gov import GrantsGovIngester

    ingester = GrantsGovIngester(db)
    try:
        if request.categories:
            result = await ingester.ingest_by_category(
                categories=request.categories,
                max_results=request.max_results,
            )
        else:
            result = await ingester.ingest_all(
                keywords=request.keywords,
                max_results=request.max_results,
            )

        return IngestResponse(
            total_fetched=result.total_fetched,
            new_grants=result.new_grants,
            updated_grants=result.updated_grants,
            details_fetched=result.details_fetched,
            scored=result.scored,
            errors=result.errors,
        )
    finally:
        await ingester.close()


class GrantExecIngestRequest(BaseModel):
    """Request to trigger GrantExec ingestion."""
    max_pages: int = 10
    fetch_details: bool = False


class GrantExecIngestResponse(BaseModel):
    """Response from GrantExec ingestion."""
    total_fetched: int
    new_grants: int
    updated_grants: int
    errors: list[str]


@app.post("/api/v1/ingest/grantexec", response_model=GrantExecIngestResponse)
async def ingest_from_grantexec(
    request: GrantExecIngestRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Ingest grants from GrantExec portal.

    Uses browser automation to scrape state/local grant opportunities.
    Requires GRANTEXEC_USERNAME and GRANTEXEC_PASSWORD environment variables.
    """
    from app.ingest.grantexec import GrantExecIngester

    ingester = GrantExecIngester(db)
    try:
        result = await ingester.ingest_all(
            max_pages=request.max_pages,
            fetch_details=request.fetch_details,
        )

        return GrantExecIngestResponse(
            total_fetched=result.total_fetched,
            new_grants=result.new_grants,
            updated_grants=result.updated_grants,
            errors=result.errors,
        )
    except ValueError as e:
        raise HTTPException(500, str(e))
    except Exception as e:
        logger.exception(f"GrantExec ingestion failed: {e}")
        raise HTTPException(500, f"Ingestion failed: {e}")


# ============================================================================
# Organization Profile
# ============================================================================

class ProfileUpdateRequest(BaseModel):
    """Request to update the organization profile."""
    content: str
    title: Optional[str] = None
    summary: Optional[str] = None
    changed_by: Optional[str] = None
    change_notes: Optional[str] = None


class ProfileResponse(BaseModel):
    """Organization profile response."""
    id: str
    version: int
    title: str
    content: str
    summary: Optional[str]
    is_active: bool
    changed_by: Optional[str]
    change_notes: Optional[str]
    relevance_rules: Optional[dict]
    rules_generated_at: Optional[str]
    created_at: str
    updated_at: str

    class Config:
        from_attributes = True


class ProfileVersionSummary(BaseModel):
    """Summary of a profile version."""
    id: str
    version: int
    title: str
    summary: Optional[str]
    changed_by: Optional[str]
    change_notes: Optional[str]
    is_active: bool
    created_at: str


@app.get("/api/v1/org/profile")
async def get_org_profile(db: AsyncSession = Depends(get_db)):
    """Get the current active organization profile."""
    from app.models import OrgProfile

    result = await db.execute(
        select(OrgProfile)
        .where(OrgProfile.is_active == True)
        .order_by(OrgProfile.version.desc())
        .limit(1)
    )
    profile = result.scalar_one_or_none()

    if not profile:
        raise HTTPException(404, "No organization profile found. Please create one.")

    return {
        "id": str(profile.id),
        "version": profile.version,
        "title": profile.title,
        "content": profile.content,
        "summary": profile.summary,
        "is_active": profile.is_active,
        "changed_by": profile.changed_by,
        "change_notes": profile.change_notes,
        "relevance_rules": profile.relevance_rules,
        "rules_generated_at": profile.rules_generated_at.isoformat() if profile.rules_generated_at else None,
        "created_at": profile.created_at.isoformat(),
        "updated_at": profile.updated_at.isoformat(),
    }


@app.put("/api/v1/org/profile")
async def update_org_profile(
    request: ProfileUpdateRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Update the organization profile.

    Creates a new version and marks previous versions as inactive.
    """
    from app.models import OrgProfile

    # Get current max version
    result = await db.execute(
        select(func.max(OrgProfile.version))
    )
    max_version = result.scalar() or 0

    # Deactivate all previous versions
    await db.execute(
        select(OrgProfile).where(OrgProfile.is_active == True)
    )
    # Using raw SQL for update
    from sqlalchemy import update
    await db.execute(
        update(OrgProfile).where(OrgProfile.is_active == True).values(is_active=False)
    )

    # Create new version
    new_profile = OrgProfile(
        version=max_version + 1,
        is_active=True,
        content=request.content,
        title=request.title or "Nexus Center for IDD Care Organization Profile",
        summary=request.summary,
        changed_by=request.changed_by,
        change_notes=request.change_notes,
    )
    db.add(new_profile)
    await db.commit()
    await db.refresh(new_profile)

    return {
        "id": str(new_profile.id),
        "version": new_profile.version,
        "status": "created",
        "message": f"Profile updated to version {new_profile.version}",
    }


@app.get("/api/v1/org/profile/versions")
async def list_profile_versions(
    limit: int = Query(20, le=100),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """List all profile versions."""
    from app.models import OrgProfile

    result = await db.execute(
        select(OrgProfile)
        .order_by(OrgProfile.version.desc())
        .offset(offset)
        .limit(limit)
    )
    profiles = result.scalars().all()

    return {
        "versions": [
            {
                "id": str(p.id),
                "version": p.version,
                "title": p.title,
                "summary": p.summary,
                "changed_by": p.changed_by,
                "change_notes": p.change_notes,
                "is_active": p.is_active,
                "created_at": p.created_at.isoformat(),
            }
            for p in profiles
        ],
        "total": len(profiles),
    }


@app.get("/api/v1/org/profile/versions/{version}")
async def get_profile_version(
    version: int,
    db: AsyncSession = Depends(get_db),
):
    """Get a specific profile version."""
    from app.models import OrgProfile

    result = await db.execute(
        select(OrgProfile).where(OrgProfile.version == version)
    )
    profile = result.scalar_one_or_none()

    if not profile:
        raise HTTPException(404, f"Profile version {version} not found")

    return {
        "id": str(profile.id),
        "version": profile.version,
        "title": profile.title,
        "content": profile.content,
        "summary": profile.summary,
        "is_active": profile.is_active,
        "changed_by": profile.changed_by,
        "change_notes": profile.change_notes,
        "relevance_rules": profile.relevance_rules,
        "rules_generated_at": profile.rules_generated_at.isoformat() if profile.rules_generated_at else None,
        "created_at": profile.created_at.isoformat(),
        "updated_at": profile.updated_at.isoformat(),
    }


@app.post("/api/v1/org/profile/upload")
async def upload_org_profile(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    changed_by: str = Form(None),
    change_notes: str = Form(None),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a Word document (.docx) as the organization profile.

    - Extracts text from the document
    - Compares with current profile
    - Creates new version only if content differs
    - Auto-generates relevance scoring rules in background
    """
    from app.models import OrgProfile
    from sqlalchemy import update
    import io

    # Validate file type
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Only .docx files are supported")

    # Read and parse the Word document
    try:
        from docx import Document
        content = await file.read()
        doc = Document(io.BytesIO(content))

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure with markdown
                style_name = para.style.name if para.style else None
                if style_name and style_name.startswith('Heading'):
                    level = 1
                    try:
                        level = int(style_name.replace('Heading ', ''))
                    except:
                        pass
                    paragraphs.append(f"{'#' * level} {text}")
                else:
                    paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        extracted_content = '\n\n'.join(paragraphs)

        if not extracted_content.strip():
            raise HTTPException(400, "Document appears to be empty")

    except ImportError:
        raise HTTPException(500, "python-docx not installed")
    except Exception as e:
        logger.exception(f"Error parsing document: {e}")
        raise HTTPException(400, f"Could not parse document: {str(e)}")

    # Get current active profile
    result = await db.execute(
        select(OrgProfile)
        .where(OrgProfile.is_active == True)
        .order_by(OrgProfile.version.desc())
        .limit(1)
    )
    current_profile = result.scalar_one_or_none()

    # Check if content has changed
    if current_profile and current_profile.content.strip() == extracted_content.strip():
        return {
            "status": "unchanged",
            "message": "Document content is identical to current profile",
            "version": current_profile.version,
            "id": str(current_profile.id),
        }

    # Get max version number
    result = await db.execute(select(func.max(OrgProfile.version)))
    max_version = result.scalar() or 0

    # Deactivate previous versions
    await db.execute(
        update(OrgProfile).where(OrgProfile.is_active == True).values(is_active=False)
    )

    # Extract title from first heading or use filename
    title = "Nexus Center for IDD Care Organization Profile"
    for para in extracted_content.split('\n'):
        if para.startswith('# '):
            title = para[2:].strip()
            break

    # Create new version
    new_profile = OrgProfile(
        version=max_version + 1,
        is_active=True,
        content=extracted_content,
        title=title,
        summary=f"Uploaded from {file.filename}",
        changed_by=changed_by or "upload",
        change_notes=change_notes or f"Uploaded document: {file.filename}",
    )
    db.add(new_profile)
    await db.commit()
    await db.refresh(new_profile)

    # Schedule background task to generate relevance rules
    async def generate_rules_task():
        from app.scoring.relevance import RelevanceScorer
        from app.db import get_session

        async with get_session() as session:
            scorer = RelevanceScorer(session)
            try:
                await scorer.generate_rules_from_profile()
                logger.info(f"Generated relevance rules for profile version {new_profile.version}")
            except Exception as e:
                logger.exception(f"Failed to generate rules: {e}")
            finally:
                await scorer.close()

    background_tasks.add_task(generate_rules_task)

    return {
        "status": "created",
        "message": f"Profile updated to version {new_profile.version}. Relevance rules are being generated.",
        "version": new_profile.version,
        "id": str(new_profile.id),
        "title": new_profile.title,
        "content_preview": extracted_content[:500] + "..." if len(extracted_content) > 500 else extracted_content,
    }


# ============================================================================
# Relevance Scoring
# ============================================================================

@app.post("/api/v1/scoring/generate-rules")
async def generate_scoring_rules(db: AsyncSession = Depends(get_db)):
    """
    Generate relevance scoring rules from the current org profile.

    Uses Cortex AI to analyze the profile and extract keywords,
    categories, and other scoring criteria.
    """
    from app.scoring.relevance import RelevanceScorer

    scorer = RelevanceScorer(db)
    try:
        rules = await scorer.generate_rules_from_profile()

        if not rules:
            raise HTTPException(500, "Failed to generate rules - check logs")

        return {
            "status": "generated",
            "rules": rules.to_dict(),
            "high_priority_count": len(rules.high_priority_keywords),
            "medium_priority_count": len(rules.medium_priority_keywords),
            "low_priority_count": len(rules.low_priority_keywords),
        }
    finally:
        await scorer.close()


@app.post("/api/v1/scoring/score-all")
async def score_all_grants(db: AsyncSession = Depends(get_db)):
    """
    Score all grants in the database using current relevance rules.

    This may take a while for large numbers of grants.
    """
    from app.scoring.relevance import RelevanceScorer

    scorer = RelevanceScorer(db)
    try:
        result = await scorer.score_all_grants()

        if "error" in result:
            raise HTTPException(400, result["error"])

        return result
    finally:
        await scorer.close()


@app.post("/api/v1/grants/{grant_id}/score")
async def score_single_grant(
    grant_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Score a single grant for relevance."""
    from app.scoring.relevance import RelevanceScorer

    scorer = RelevanceScorer(db)
    try:
        result = await scorer.score_single_grant(grant_id)

        if "error" in result:
            raise HTTPException(400, result["error"])

        return result
    finally:
        await scorer.close()


@app.get("/api/v1/scoring/rules")
async def get_scoring_rules(db: AsyncSession = Depends(get_db)):
    """Get the current relevance scoring rules."""
    from app.models import OrgProfile

    result = await db.execute(
        select(OrgProfile)
        .where(OrgProfile.is_active == True)
        .order_by(OrgProfile.version.desc())
        .limit(1)
    )
    profile = result.scalar_one_or_none()

    if not profile:
        raise HTTPException(404, "No active profile found")

    if not profile.relevance_rules:
        raise HTTPException(404, "No relevance rules generated yet")

    return {
        "rules": profile.relevance_rules,
        "generated_at": profile.rules_generated_at.isoformat() if profile.rules_generated_at else None,
        "profile_version": profile.version,
    }


class RulesUpdateRequest(BaseModel):
    """Request to update relevance scoring rules."""
    high_priority_keywords: Optional[list[str]] = None
    medium_priority_keywords: Optional[list[str]] = None
    low_priority_keywords: Optional[list[str]] = None
    negative_keywords: Optional[list[str]] = None
    relevant_categories: Optional[list[str]] = None
    relevant_agencies: Optional[list[str]] = None
    geographic_keywords: Optional[list[str]] = None
    population_keywords: Optional[list[str]] = None
    min_preferred_funding: Optional[float] = None
    max_preferred_funding: Optional[float] = None


@app.patch("/api/v1/scoring/rules")
async def update_scoring_rules(
    request: RulesUpdateRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Update relevance scoring rules.

    Only updates the fields that are provided. Pass an empty list to clear a field.
    """
    from app.models import OrgProfile
    from datetime import datetime as dt, timezone

    result = await db.execute(
        select(OrgProfile)
        .where(OrgProfile.is_active == True)
        .order_by(OrgProfile.version.desc())
        .limit(1)
    )
    profile = result.scalar_one_or_none()

    if not profile:
        raise HTTPException(404, "No active profile found")

    if not profile.relevance_rules:
        raise HTTPException(400, "No relevance rules to update. Generate rules first.")

    # Update only provided fields
    rules = dict(profile.relevance_rules)

    if request.high_priority_keywords is not None:
        rules["high_priority_keywords"] = request.high_priority_keywords
    if request.medium_priority_keywords is not None:
        rules["medium_priority_keywords"] = request.medium_priority_keywords
    if request.low_priority_keywords is not None:
        rules["low_priority_keywords"] = request.low_priority_keywords
    if request.negative_keywords is not None:
        rules["negative_keywords"] = request.negative_keywords
    if request.relevant_categories is not None:
        rules["relevant_categories"] = request.relevant_categories
    if request.relevant_agencies is not None:
        rules["relevant_agencies"] = request.relevant_agencies
    if request.geographic_keywords is not None:
        rules["geographic_keywords"] = request.geographic_keywords
    if request.population_keywords is not None:
        rules["population_keywords"] = request.population_keywords
    if request.min_preferred_funding is not None:
        rules["min_preferred_funding"] = request.min_preferred_funding
    if request.max_preferred_funding is not None:
        rules["max_preferred_funding"] = request.max_preferred_funding

    profile.relevance_rules = rules
    profile.rules_generated_at = dt.now(timezone.utc)

    await db.commit()

    return {
        "status": "updated",
        "rules": rules,
        "updated_at": profile.rules_generated_at.isoformat(),
    }


# ============================================================================
# Static file serving (frontend) — must be last
# ============================================================================

# Mount static files if the frontend dist directory exists
import pathlib
static_dir = pathlib.Path(__file__).parent.parent / "frontend" / "dist"
if static_dir.exists():
    from starlette.responses import FileResponse

    @app.get("/{full_path:path}")
    async def serve_frontend(full_path: str):
        """Serve frontend static files, falling back to index.html for SPA routing."""
        file_path = static_dir / full_path
        if file_path.exists() and file_path.is_file():
            return FileResponse(file_path)
        return FileResponse(static_dir / "index.html")
